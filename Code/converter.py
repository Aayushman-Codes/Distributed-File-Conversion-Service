"""
converter.py — File conversion logic for the DFS.

Supported conversions:
  Images  (Pillow):        jpg/jpeg/png/bmp/gif/webp/tiff  ↔ any of the same
  Text    (stdlib):        txt / csv / json / xml           ↔ any of the same
  Docs    (pypdf/docx):    pdf → txt/docx,  docx → txt/pdf

Returns the MD5 hex-digest of the output file.
"""

import os
import csv
import json
import hashlib
import logging
import xml.etree.ElementTree as ET
from io import StringIO

logger = logging.getLogger("converter")

IMAGE_FORMATS = {"jpg", "jpeg", "png", "bmp", "gif", "webp", "tiff"}
TEXT_FORMATS  = {"txt", "csv", "json", "xml"}
DOC_FORMATS   = {"pdf", "docx"}

PIL_FORMAT_MAP = {
    "jpg": "JPEG", "jpeg": "JPEG", "png": "PNG",
    "bmp": "BMP",  "gif": "GIF",  "webp": "WEBP", "tiff": "TIFF",
}


def convert_file(in_path, out_path, src_fmt, dst_fmt):
    src_fmt = src_fmt.lower()
    dst_fmt = dst_fmt.lower()

    if src_fmt in IMAGE_FORMATS and dst_fmt in IMAGE_FORMATS:
        _convert_image(in_path, out_path, dst_fmt)
    elif src_fmt in TEXT_FORMATS and dst_fmt in TEXT_FORMATS:
        _convert_text(in_path, out_path, src_fmt, dst_fmt)
    elif src_fmt in DOC_FORMATS or dst_fmt in DOC_FORMATS:
        _convert_document(in_path, out_path, src_fmt, dst_fmt)
    else:
        raise ValueError(
            f"Unsupported conversion: {src_fmt} → {dst_fmt}. "
            f"Images: {sorted(IMAGE_FORMATS)}, "
            f"Text: {sorted(TEXT_FORMATS)}, "
            f"Docs: {sorted(DOC_FORMATS)}"
        )
    return _md5_file(out_path)


# ── Image ─────────────────────────────────────────────────────────────────────

def _convert_image(in_path, out_path, dst_fmt):
    from PIL import Image
    pil_fmt = PIL_FORMAT_MAP[dst_fmt]
    with Image.open(in_path) as img:
        if pil_fmt == "JPEG" and img.mode in ("RGBA", "P", "LA"):
            img = img.convert("RGB")
        img.save(out_path, format=pil_fmt, quality=95)


# ── Text / Data ───────────────────────────────────────────────────────────────

def _convert_text(in_path, out_path, src_fmt, dst_fmt):
    with open(in_path, "r", encoding="utf-8") as f:
        raw = f.read()

    # Parse
    if src_fmt == "json":
        data = json.loads(raw)
    elif src_fmt == "csv":
        data = list(csv.DictReader(StringIO(raw)))
    elif src_fmt == "xml":
        data = _xml_to_dict(raw)
    else:  # txt
        data = raw.splitlines()

    # Serialize
    if dst_fmt == "json":
        out = json.dumps(data, indent=2, ensure_ascii=False)
    elif dst_fmt == "csv":
        out = _to_csv_string(data)
    elif dst_fmt == "xml":
        out = _to_xml_string(data)
    else:  # txt
        out = _to_txt_string(data)

    with open(out_path, "w", encoding="utf-8") as f:
        f.write(out)


def _xml_to_dict(raw):
    """Parse XML into a list of dicts (one per child of root)."""
    root = ET.fromstring(raw)
    result = []
    for child in root:
        row = {"_tag": child.tag}
        row.update(child.attrib)
        for sub in child:
            row[sub.tag] = sub.text or ""
        if not list(child):
            row["_text"] = child.text or ""
        result.append(row)
    return result if result else [{"_tag": root.tag, "_text": root.text or ""}]


def _to_xml_string(data):
    root = ET.Element("root")
    if isinstance(data, list):
        for item in data:
            if isinstance(item, dict):
                tag = item.get("_tag", "item")
                child = ET.SubElement(root, tag)
                for k, v in item.items():
                    if k not in ("_tag", "_text"):
                        sub = ET.SubElement(child, k)
                        sub.text = str(v)
                if "_text" in item:
                    child.text = item["_text"]
            else:
                child = ET.SubElement(root, "item")
                child.text = str(item)
    else:
        child = ET.SubElement(root, "data")
        child.text = str(data)
    ET.indent(root)
    return ET.tostring(root, encoding="unicode", xml_declaration=False)


def _to_csv_string(data):
    if not data:
        return ""
    buf = StringIO()
    if isinstance(data, list) and isinstance(data[0], dict):
        writer = csv.DictWriter(buf, fieldnames=data[0].keys(),
                                extrasaction="ignore")
        writer.writeheader()
        writer.writerows(data)
    else:
        writer = csv.writer(buf)
        for item in data:
            writer.writerow([item])
    return buf.getvalue()


def _to_txt_string(data):
    if isinstance(data, list):
        return "\n".join(str(x) for x in data)
    return str(data)


# ── Documents ─────────────────────────────────────────────────────────────────

def _convert_document(in_path, out_path, src_fmt, dst_fmt):
    if src_fmt == "pdf" and dst_fmt == "txt":
        _pdf_to_txt(in_path, out_path)
    elif src_fmt == "pdf" and dst_fmt == "docx":
        _pdf_to_docx(in_path, out_path)
    elif src_fmt == "docx" and dst_fmt == "txt":
        _docx_to_txt(in_path, out_path)
    elif src_fmt == "docx" and dst_fmt == "pdf":
        _docx_to_pdf(in_path, out_path)
    else:
        raise ValueError(f"Unsupported document conversion: {src_fmt} → {dst_fmt}")


def _pdf_to_txt(in_path, out_path):
    """Extract all text from a PDF and write to a UTF-8 text file."""
    import pypdf
    reader = pypdf.PdfReader(in_path)
    pages_text = []
    for page in reader.pages:
        text = page.extract_text(extraction_mode="layout") or ""
        pages_text.append(text.strip())
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n\n".join(t for t in pages_text if t))


def _pdf_to_docx(in_path, out_path):
    """Extract PDF text page-by-page and write into a DOCX document."""
    import pypdf
    from docx import Document as DocxDocument
    from docx.shared import Pt
    reader = pypdf.PdfReader(in_path)
    doc = DocxDocument()
    doc.add_heading("Converted from PDF", level=1)
    for i, page in enumerate(reader.pages):
        text = (page.extract_text(extraction_mode="layout") or "").strip()
        doc.add_heading(f"Page {i + 1}", level=2)
        if text:
            for line in text.splitlines():
                stripped = line.strip()
                if stripped:
                    p = doc.add_paragraph(stripped)
                    p.style.font.size = Pt(11)
        else:
            doc.add_paragraph("(no extractable text on this page)")
    doc.save(out_path)


def _docx_to_txt(in_path, out_path):
    """Extract all text from a DOCX (paragraphs + tables) into a text file."""
    from docx import Document as DocxDocument
    doc = DocxDocument(in_path)
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            lines.append("\t".join(row_cells))
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


def _docx_to_pdf(in_path, out_path):
    """Convert DOCX → PDF using reportlab, preserving headings, body text, tables and images."""
    import io
    import xml.sax.saxutils as saxutils
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
    )

    PAGE_W = A4[0] - 4 * cm   # usable width (left+right margin = 2cm each)

    def safe(t):
        return saxutils.escape(t.replace("\x00", ""))

    doc = DocxDocument(in_path)
    styles = getSampleStyleSheet()

    heading_styles = {}
    for level in range(1, 10):
        rl_name = f"Heading{level}"
        heading_styles[level] = styles.get(rl_name, styles["Heading1"])

    # Build a map: rId → image bytes  (covers inline and anchored images)
    image_map: dict[str, bytes] = {}
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                image_map[rel.rId] = rel.target_part.blob
            except Exception:
                pass

    def _para_images(para):
        """Return list of (rId, width_emu) for every inline/anchored image in a paragraph."""
        found = []
        # Inline drawings: w:drawing > wp:inline > a:graphic > ... > a:blip r:embed
        for drawing in para._element.iter(qn("w:drawing")):
            # try inline first
            for blip in drawing.iter(qn("a:blip")):
                rId = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
                if rId and rId in image_map:
                    # get width from wp:extent cx attribute (EMUs)
                    width_emu = 0
                    for extent in drawing.iter(qn("wp:extent")):
                        try:
                            width_emu = int(extent.get("cx", 0))
                        except (ValueError, TypeError):
                            pass
                    found.append((rId, width_emu))
        return found

    pdf_doc = SimpleDocTemplate(
        out_path, pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
    )

    story = []

    for para in doc.paragraphs:
        # ── Images embedded in this paragraph ────────────────────────────────
        for rId, width_emu in _para_images(para):
            img_bytes = image_map[rId]
            try:
                # Convert EMU to points (1 inch = 914400 EMU = 72 pt)
                if width_emu > 0:
                    w_pt = width_emu / 914400 * 72
                    w_pt = min(w_pt, PAGE_W)   # never wider than the page
                else:
                    w_pt = PAGE_W              # default: full width
                rl_img = RLImage(io.BytesIO(img_bytes), width=w_pt)
                rl_img.hAlign = "LEFT"
                story.append(rl_img)
                story.append(Spacer(1, 6))
            except Exception as exc:
                logger.warning("Skipping image %s: %s", rId, exc)

        # ── Text ──────────────────────────────────────────────────────────────
        text = para.text
        style_name = para.style.name or ""

        if not text.strip():
            story.append(Spacer(1, 4))
            continue

        if style_name.startswith("Heading"):
            parts = style_name.split()
            try:
                level = int(parts[-1])
            except (ValueError, IndexError):
                level = 1
            rl_style = heading_styles.get(level, styles["Heading1"])
        else:
            rl_style = styles["Normal"]

        story.append(Paragraph(safe(text), rl_style))
        story.append(Spacer(1, 4))

    # ── Tables ────────────────────────────────────────────────────────────────
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            table_data.append([safe(cell.text) for cell in row.cells])

        if not table_data:
            continue

        rl_table = Table(table_data, repeatRows=1)
        rl_table.setStyle(TableStyle([
            ("BACKGROUND",   (0, 0), (-1, 0),  colors.HexColor("#4f6ef7")),
            ("TEXTCOLOR",    (0, 0), (-1, 0),  colors.white),
            ("FONTNAME",     (0, 0), (-1, 0),  "Helvetica-Bold"),
            ("FONTSIZE",     (0, 0), (-1, -1), 9),
            ("ROWBACKGROUNDS",(0, 1),(-1, -1), [colors.white, colors.HexColor("#f0f0f8")]),
            ("GRID",         (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
            ("VALIGN",       (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING",  (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(Spacer(1, 8))
        story.append(rl_table)
        story.append(Spacer(1, 8))

    if not story:
        story.append(Paragraph("(empty document)", styles["Normal"]))

    pdf_doc.build(story)


# ── Utility ───────────────────────────────────────────────────────────────────

def _md5_file(path):
    h = hashlib.md5()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()
